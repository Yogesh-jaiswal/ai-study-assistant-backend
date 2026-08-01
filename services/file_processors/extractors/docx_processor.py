from pathlib import Path
import tempfile
import zipfile

from PIL import Image
from docx import Document

from services.integrations.ocr_service import OCR

from .base_processor import BaseProcessor

from models.enums import DocumentBlockType

from services.file_processors.document.doc_representation import (
    DocumentBlock,
    DocumentRepresentation
)

class DOCXProcessor(BaseProcessor):
    """DOCX file processor"""

    def __init__(self):
        self.ocr = OCR()

    def extract(self, file_path: str | Path) -> DocumentRepresentation:
        document = Document(file_path)
        props = document.core_properties

        blocks = []

        # Paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            block_type = (
                DocumentBlockType.HEADING
                if paragraph.style.name.startswith("Heading")
                else DocumentBlockType.PARAGRAPH
            )

            if text:
                blocks.append(
                    DocumentBlock(
                        type=block_type,
                        text=text,
                    )
                )

        # Tables
        for table in document.tables:
            rows = [
                " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                )
                for row in table.rows
            ]

            text = "\n".join(rows).strip()

            if text:
                blocks.append(
                    DocumentBlock(
                        type=DocumentBlockType.TABLE,
                        text=text,
                    )
                )

        # Images (OCR)
        blocks.extend(self._extract_images(file_path))

        return DocumentRepresentation(
            author=props.author or None,
            blocks=blocks
        )

    def _extract_images(self, file_path: str | Path) -> list[DocumentBlock]:
        with tempfile.TemporaryDirectory() as temp_dir:
            with zipfile.ZipFile(file_path) as archive:
                archive.extractall(temp_dir)

            media_dir = Path(temp_dir) / "word" / "media"

            if not media_dir.exists():
                return []

            output = []

            for image_path in sorted(media_dir.iterdir()):
                if not image_path.is_file():
                    continue

                with Image.open(image_path) as img:
                    text = self.ocr.extract_text(img)

                if text:
                    output.append(
                        DocumentBlock(
                            type=DocumentBlockType.OCR,
                            text=text,
                        )
                    )

            return output